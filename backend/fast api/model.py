import os
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
import fitz  # PyMuPDF
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.embeddings import HuggingFaceEmbeddings
from docx import Document
from langchain.schema import Document as LangchainDocument


class DocumentProcessor:
    def __init__(self, input_folder, output_folder):
        self.input_folder = input_folder
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)

    def extract_text_from_pdf(self, pdf_path):
        text = ""

        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception as e:
            print(f"Ошибка при извлечении текста через PyPDF2: {e}")

        if text.strip():
            return text

        try:
            images = convert_from_path(pdf_path, 300)
            for img in images:
                text += pytesseract.image_to_string(img)
        except Exception as e:
            print(f"Ошибка при извлечении текста через Tesseract: {e}")

        if text.strip():
            return text

        try:
            doc = fitz.open(pdf_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text("text")
        except Exception as e:
            print(f"Ошибка при извлечении текста через PyMuPDF: {e}")

        return text.strip() if text.strip() else None

    def clean_text(self, text):
        return ' '.join(text.replace('\n', ' ').split())

    def process_pdf(self, pdf_path):
        text = self.extract_text_from_pdf(pdf_path)
        if text:
            cleaned_text = self.clean_text(text)
            output_file = os.path.splitext(pdf_path)[0] + "_extracted.txt"
            with open(output_file, "w", encoding="utf-8") as file:
                file.write(cleaned_text)
            print(f"Текст успешно извлечен и сохранен в файл: {output_file}")
        else:
            print("Не удалось извлечь текст из PDF.")

    def extract_text_to_folder(self):
        pdf_files = [f for f in os.listdir(self.input_folder) if f.endswith('.pdf')]

        if not pdf_files:
            print("Нет PDF-файлов в указанной папке.")
            return

        for pdf_file in pdf_files:
            pdf_path = os.path.join(self.input_folder, pdf_file)
            output_file = os.path.join(
                self.output_folder, os.path.splitext(pdf_file)[0] + ".txt"
            )

            try:
                reader = PdfReader(pdf_path)
                with open(output_file, "w", encoding="utf-8") as f:
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = page.extract_text()
                        if text:
                            f.write(f"Page {page_num}\n{text.strip()}\n\n")
                        else:
                            f.write(f"Page {page_num}\nСтраница пуста или текст не удалось извлечь.\n\n")
                print(f"Текст извлечен из {pdf_file} и сохранен в {output_file}")
            except Exception as e:
                print(f"Ошибка при обработке {pdf_file}: {e}")

    def read_docx(self, file_path):
        doc = Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs)
        return text.strip()

    def split_text_into_fragments(self, text, fragment_length=400, overlap=100):
        start = 0
        fragments = []
        while start < len(text):
            end = min(start + fragment_length, len(text))
            fragment = text[start:end]
            fragments.append(
                LangchainDocument(page_content=fragment, metadata={})
            )
            start += fragment_length - overlap
        return fragments

    def process_docx_and_create_index(self, fragment_length=400, overlap=100):
        docx_files = [f for f in os.listdir(self.input_folder) if f.endswith('.docx')]

        if not docx_files:
            print("Нет .docx файлов в указанной папке.")
            return

        all_fragments = []
        for docx_file in docx_files:
            file_path = os.path.join(self.input_folder, docx_file)
            text = self.read_docx(file_path)
            fragments = self.split_text_into_fragments(
                text, fragment_length=fragment_length, overlap=overlap
            )
            all_fragments.extend(fragments)

        return all_fragments

    # Неиспользуемые функции:
    def save_text_to_docx(self, text, output_path):
        """Сохраняет текст в формате .docx."""
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        print(f"Текст сохранен в {output_path}")

    def dummy_method(self):
        """Просто метод-заглушка."""
        print("Это метод-заглушка. Он ничего не делает.")

    def unused_ocr_on_image(self, image_path):
        """OCR обработка изображения (не используется)."""
        try:
            text = pytesseract.image_to_string(image_path)
            print(f"Текст на изображении:\n{text}")
            return text
        except Exception as e:
            print(f"Ошибка OCR обработки изображения: {e}")
            return None

    def get_embedding_from_gigachain(self, text):
        """Неиспользуемая функция для получения эмбединга."""
        context = (
            "Во время предоставления представь, что ты составляешь эмбединг..."
        )

        prompt_template = PromptTemplate(
            input_variables=["text", "context"],
            template="{context}\n\nТекст: {text}",
        )

        llm = GigaChat(
            credentials="ВАШ_AUTH",
            model="GigaChat:latest",
            verify_ssl_certs=False,
            profanity_check=False,
        )

        chain = LLMChain(llm=llm, prompt=prompt_template)

        try:
            return chain.run({"text": text, "context": context})
        except Exception as e:
            print(f"Ошибка при выполнении запроса через GigaChain: {e}")
            return None

    from torch.utils.data import DataLoader
import torch
from PIL import Image
import os
from colpali_engine.models import ColQwen2, ColQwen2Processor
import gc

# Загрузка модели и процессора для пост обработки нужных эмбедингов
model_name = "vidore/colqwen2-v0.1"
model = ColQwen2.from_pretrained(
    model_name,
    torch_dtype=torch.float16,  # Используем половинную точность
    device_map="cuda:0",
).eval()
processor = ColQwen2Processor.from_pretrained(model_name)

# Функция для обработки изображений по батчам
def process_images_in_batches(model, processor, images, batch_size, query):

    # Путь к папке с изображениями
    image_folder = "output_images"
    images = [Image.open(os.path.join(image_folder, img)) for img in os.listdir(image_folder)]

    # Обработка текста
    batch_query = processor.process_queries([query]).to(model.device, dtype=torch.float16)
    embeddings = []
    dataloader = DataLoader(images, batch_size=batch_size, shuffle=False)

    with torch.no_grad():
        for batch in dataloader:
            batch_images = processor.process_images(batch).to(model.device, dtype=torch.float16)
            batch_embeddings = model(**batch_images)
            embeddings.append(batch_embeddings)

            # Очистка памяти
            del batch_images, batch_embeddings
            torch.cuda.empty_cache()
            gc.collect()

    return torch.cat(embeddings, dim=0)

    from docx import Document
from langchain.schema import Document as LangchainDocument

def read_docx(file_path):
    """Извлекает текст из документа .docx"""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text.strip()

def split_text_into_fragments(text, fragment_length=400, overlap=100):
    """
    Разделяет текст на фрагменты с наложением.
    """
    start = 0
    fragments_with_metadata = []
    while start < len(text):
        end = min(start + fragment_length, len(text))
        fragment = text[start:end]
        fragments_with_metadata.append(
            LangchainDocument(page_content=fragment, metadata={})
        )
        start += fragment_length - overlap  # Двигаемся с учетом наложения
    return fragments_with_metadata

def process_docx_and_create_index(input_folder, output_index_file, fragment_length=400, overlap=100):
    """Обрабатывает все .docx файлы в папке и создает общий векторный индекс."""
    all_fragments = []
    docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]

    if not docx_files:
        print("Нет .docx файлов в указанной папке.")
        return

    for docx_file in docx_files:
        file_path = os.path.join(input_folder, docx_file)
        print(f"Обрабатывается файл: {docx_file}")
        text = read_docx(file_path)
        fragments = split_text_into_fragments(text, fragment_length=fragment_length, overlap=overlap)
        all_fragments.extend(fragments)

    # Параметры модели
    embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")

    # Создание векторного хранилища
    vector_store = FAISS.from_documents(all_fragments, embedding)

    # Сохранение индекса
    vector_store.save_local(output_index_file)
    print(f"Индекс с метаданными сохранен в файл: {output_index_file}")


from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate

def get_embedding_from_gigachain(text):
    # Определение контекста для обработки текста
    context = (
        "Во время предоставления представь что ты составляешь эмбединг, который потом будет прогоняться через вектор, ответ не должен превышать 500 символов"
        "Можешь переписать текст, так чтобы выделить в нем главную информацию и найди в этом "
        "фрагменте значения полей графиков и преобразуй их в таблицу по твоему усмотрению.\n"
        "Необходима таблица вида например:\n"
        "+-------------------+-------------------+\n"
        "|      Заголовок 1  |      Заголовок 2  |\n"
        "+-------------------+-------------------+\n"
        "|      Значение 1   |      Значение 2   |\n"
        "+-------------------+-------------------+\n"
        "|      Значение 3   |      Значение 4   |\n"
        "+-------------------+-------------------+\n"
        "или вида:\n"
        "+-------------------+-------------------+\n"
        "| Заголовок 1       | Заголовок 2       |\n"
        "+-------------------+-------------------+\n"
        "| Верхняя строка 1  | Верхняя строка 2  |\n"
        "+-------------------+-------------------+\n"
        "| Значение 1        | Значение 2        |\n"
        "| Значение 3        | Значение 4        |\n"
        "+-------------------+-------------------+\n"
        "Ну или любая другая таблица на твое усмотрение, лишь бы ты понял."
    )

    # Настраиваем шаблон запроса
    prompt_template = PromptTemplate(
        input_variables=["text", "context"],
        template="{context}\n\nТекст: {text}"
    )

    auth = "MTA4ZDM1OTQtODA5Yi00NDMzLTk3Y2ItZjNjNTZiODRkOGY0Ojk5YjI5MzUwLThkYzYtNGJhMy04ZGQ5LTk5MTYwYmNhYmRmMg=="

    llm = GigaChat(
        credentials=auth,
        model="GigaChat:latest",  # Убедитесь, что эта модель доступна
        verify_ssl_certs=False,
        profanity_check=False
    )

    # Создаем цепочку для выполнения запроса
    chain = LLMChain(
        llm=llm,
        prompt=prompt_template
    )

    # Выполняем запрос и обрабатываем результат
    try:
        result = chain.run({"text": text, "context": context})
        return result  # Результат должен включать таблицу или обработанный текст
    except Exception as e:
        print(f"Ошибка при выполнении запроса через GigaChain: {e}")
        return None

    def extract_text_from_pdf(pdf_path):
        """Извлекает текст из PDF с использованием разных методов."""
        text = ""

        try:
            # Попробуем использовать PyPDF2 для извлечения текста
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception as e:
            print(f"Ошибка при извлечении текста через PyPDF2: {e}")

        if text.strip():
            return text

        try:
            # Если PyPDF2 не смог извлечь текст, попробуем использовать pdf2image и Tesseract
            images = convert_from_path(pdf_path, 300)  # Конвертируем PDF в изображения
            text = ""
            for img in images:
                text += pytesseract.image_to_string(img)  # Преобразуем изображение в текст
        except Exception as e:
            print(f"Ошибка при извлечении текста через Tesseract: {e}")

        if text.strip():
            return text

        try:
            # Используем PyMuPDF (fitz) для извлечения текста
            doc = fitz.open(pdf_path)
            text = ""
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text("text")  # Получаем текст в обычном формате
        except Exception as e:
            print(f"Ошибка при извлечении текста через PyMuPDF: {e}")

        return text.strip() if text.strip() else None

    def clean_text(text):
        """Очистка текста от лишних пробелов и символов."""
        text = text.replace('\n', ' ')  # Заменяем переводы строк на пробелы
        text = ' '.join(text.split())  # Убираем лишние пробелы

        return result

    def process_pdf(pdf_path):
        """Основная функция для обработки PDF и восстановления текста."""
        text = extract_text_from_pdf(pdf_path)
        if text:
            cleaned_text = clean_text(text)
            output_file = os.path.splitext(pdf_path)[0] + "_extracted.txt"
            with open(output_file, "w", encoding="utf-8") as file:
                file.write(cleaned_text)
            print(f"Текст успешно извлечен и сохранен в файл: {output_file}")
        else:
            print("Не удалось извлечь текст из PDF.")



